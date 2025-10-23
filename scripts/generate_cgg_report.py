# -*- coding: utf-8 -*-
import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Optional lightweight PDF generation (text-focused)
try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

DOWNLOADS = str(Path.home() / "Downloads")
OUT_DOCX = os.path.join(DOWNLOADS, "Rapport_final_-_Projet_Synthese_groupe_3503_CGG_v1.2.docx")
OUT_PDF  = os.path.join(DOWNLOADS, "Rapport_final_-_Projet_Synthese_groupe_3503_v1.2.pdf")

BLUE = RGBColor(0x1F, 0x4E, 0x79)  # CGG-like blue for headings


def set_styles(doc: Document):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for level in (1, 2, 3, 4):
        try:
            s = doc.styles[f'Heading {level}']
        except KeyError:
            continue
        s.font.name = 'Times New Roman'
        s.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        s.font.color.rgb = BLUE


def add_field_run(paragraph, instr_text: str):
    # Creates a complex field like PAGE or NUMPAGES
    run = paragraph.add_run()
    r_element = run._r
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr_text

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    return run


def add_footer_page_numbers(doc: Document):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Page ")
    add_field_run(p, 'PAGE')
    p.add_run(" / ")
    add_field_run(p, 'NUMPAGES')


def add_cover_page(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    # Centered title block
    spacer = doc.add_paragraph("\n\n\n\n")
    spacer.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("RAPPORT FINAL – PROJET SYNTHÈSE\n\n")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = BLUE

    app = doc.add_paragraph()
    app.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app.add_run("Application: InkMatching\n\n").font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Programme: Informatique – Projet de Synthèse\n").font.size = Pt(12)
    meta.add_run("Groupe: 3503\n").font.size = Pt(12)
    meta.add_run("Date de remise: 21 octobre 2025\n").font.size = Pt(12)
    meta.add_run("Auteur: Youness Haji\n\n").font.size = Pt(12)

    college = doc.add_paragraph()
    college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = college.add_run("Collège Gérald-Godin, Montréal (QC)")
    rr.font.size = Pt(12)
    rr.bold = True

    # Page break to start content
    doc.add_page_break()


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Table des matières (clic droit > Mettre à jour le champ dans Word)")
    r.italic = True
    # Insert a TOC field (Word must update to render)
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = ""
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    doc.add_page_break()


def add_table(doc: Document, headers, rows, col_widths_pt=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths_pt:
        for row in table.rows:
            for i, w in enumerate(col_widths_pt):
                row.cells[i].width = Inches(w / 72.0)
    doc.add_paragraph("\n")


def add_code_block(doc: Document, text: str):
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)


def add_content(doc: Document):
    # HISTORIQUE DES VERSIONS
    doc.add_heading('HISTORIQUE DES VERSIONS', level=1)
    add_table(
        doc,
        headers=["Version", "Date", "Auteur", "Description"],
        rows=[
            ["1.0", "2025-10-21", "Y. Haji", "Version initiale (PDF)"],
            ["1.1", "2025-10-21", "Y. Haji", "Mise en forme complète (DOCX): TOC, figures, tableaux, acronymes, en-têtes/pieds"],
            ["1.2", "2025-10-21", "Y. Haji", "Complétion selon gabarit professeur"],
        ],
    )

    # LISTE DES FIGURES / TABLEAUX / ACRONYMES
    doc.add_heading('TABLE DES MATIÈRES', level=1)
    doc.add_paragraph('(Voir TOC ci-dessus, à mettre à jour dans Word)')

    doc.add_heading('LISTE DES FIGURES', level=1)
    figs = [
        "Figure 1 – Diagramme de contexte",
        "Figure 2 – Diagramme des cas d'utilisation",
        "Figure 3 – Diagramme de classes (modèle de données)",
        "Figure 4 – Diagramme de séquence – S'enregistrer",
        "Figure 5 – Diagramme de séquence – Recherche d'artistes",
        "Figure 6 – Diagramme de séquence – Client envoie un message",
        "Figure 7 – Diagramme de séquence – Artiste répond",
        "Figure 8 – Diagramme de séquence – Client consulte les stencils",
        "Figure 9 – Diagramme de transition d'états – Booking",
        "Figure 10 – Diagramme de transition d'états – Général",
        "Figure 11 – Architecture matérielle",
        "Figure 12 – Architecture logicielle",
    ]
    for f in figs:
        doc.add_paragraph(f, style=None)

    doc.add_heading('LISTE DES TABLEAUX', level=1)
    tabs = [
        "Tableau 1 – Historique des versions",
        "Tableau 2 – Acronymes et abréviations",
        "Tableau 3 – Structure de la base de données",
        "Tableau 4 – Exigences logiques de base de données",
        "Tableau 5 – Exigences non-fonctionnelles",
    ]
    for t in tabs:
        doc.add_paragraph(t)

    doc.add_heading('LISTE DES ACRONYMES ET ABRÉVIATIONS', level=1)
    add_table(
        doc,
        headers=["Acronyme", "Définition"],
        rows=[
            ["RTDB", "Firebase Realtime Database"],
            ["CRUD", "Create, Read, Update, Delete"],
            ["UI", "User Interface"],
            ["UX", "User Experience"],
            ["SDK", "Software Development Kit"],
            ["TOC", "Table Of Contents"],
            ["API", "Application Programming Interface"],
            ["JWT", "JSON Web Token (auth côté web si utilisé)"],
            ["SaaS", "Software as a Service"],
            ["SLA", "Service Level Agreement"],
        ],
    )

    doc.add_paragraph("Note : Ce document vise à présenter le système logiciel dans son ensemble et à établir les spécifications et exigences qui seront validées à la livraison du système. Il doit être présenté de la façon la plus directe et concise possible en évitant les répétitions.")

    # 1. INTRODUCTION
    doc.add_heading('1. INTRODUCTION', level=1)
    doc.add_paragraph(
        "InkMatching est une application web et mobile conçue pour faciliter la mise en relation entre clients et artistes tatoueurs. "
        "Le système offre une plateforme complète permettant la découverte d'artistes, la gestion des demandes de tatouage (leads), la communication en temps réel, la réservation de rendez-vous et le suivi des soins post-tatouage (aftercare)."
    )
    doc.add_paragraph(
        "Ce document présente la vue d'ensemble du système logiciel InkMatching, incluant ses spécifications fonctionnelles et non-fonctionnelles, son architecture technique, ainsi que les exigences qui seront validées à la livraison du système."
    )

    doc.add_heading('1.1 Objectifs', level=2)
    for bullet in [
        "Documenter l'ensemble des exigences fonctionnelles et non-fonctionnelles du système InkMatching",
        "Présenter l'architecture technique et les choix technologiques",
        "Décrire les spécifications détaillées permettant la conception et les tests du système",
        "Servir de référence pour l'équipe de développement et les évaluateurs du projet",
        "Établir les critères de validation et d'acceptation du système",
    ]:
        p = doc.add_paragraph(style=None)
        p.add_run("• ").bold = True
        p.add_run(bullet)
    doc.add_paragraph("Public-cible: équipe de développement, professeurs/évaluateurs CGG, parties prenantes.")

    doc.add_heading('1.2 Portée', level=2)
    doc.add_paragraph("Nom du système : InkMatching")
    doc.add_paragraph(
        "Système biplateforme (web et iOS) connectant des clients et des artistes. Backend Firebase. Fonctionnalités incluses : profils publics (adresse, ville, styles, photo, lat/lng), carte, géocodage (MapLibre + Nominatim), leads, messagerie temps réel, calendrier/booking, aftercare, stencils, authentification. Hors portée : paiements avancés, notation, vocal/vidéo, galerie publique, recommandations automatiques."
    )

    doc.add_heading('1.3 Définitions', level=2)
    defs = [
        ("Lead", "Demande initiale d'un client auprès d'un artiste."),
        ("Aftercare", "Plan de soins post-tatouage assigné au client."),
        ("Profil public", "Fiche visible: displayName, adresse, ville, styles, cover, lat/lng."),
        ("Thread", "Conversation privée client-artiste."),
        ("Stencil", "Modèle/dessin préparatoire uploadé par l'artiste."),
        ("Géocodage", "Conversion d'adresse en coordonnées GPS."),
        ("Status (lead)", "pending, accepted, rejected, completed."),
        ("Booking", "Réservation rendez-vous avec date/heure/durée/statut."),
        ("RTDB", "Base NoSQL temps réel de Firebase."),
        ("UID", "Identifiant unique Firebase Auth."),
    ]
    for k, v in defs:
        doc.add_paragraph(f"• {k} : {v}")

    doc.add_heading('1.4 Documents de références', level=2)
    refs = [
        "Code source (web Next.js/TS, iOS SwiftUI)",
        "Firebase Docs", "Next.js Docs", "SwiftUI Docs", "MapLibre GL JS Docs", "Nominatim API",
        "Diagrammes UML (annexe)",
        "Gabarit de rapport final – professeur",
        "Normes: TS/Swift, règles Firebase, Material/HIG",
    ]
    for r in refs:
        doc.add_paragraph(f"- {r}")

    # 2. DESCRIPTION GÉNÉRALE DU LOGICIEL
    doc.add_heading('2. DESCRIPTION GÉNÉRALE DU LOGICIEL', level=1)
    doc.add_heading("2.1 Vue d'ensemble des fonctions du produit", level=2)
    doc.add_paragraph("Découverte: profils publics, carte interactive, filtres.")
    doc.add_paragraph("Communication: leads, messagerie temps réel, notifications.")
    doc.add_paragraph("Gestion: calendrier/booking, aftercare, stencils, profil.")

    doc.add_heading('2.2 Interfaces externes', level=2)
    doc.add_paragraph("Acteurs: Client, Artiste. Systèmes: Firebase Auth/RTDB/Storage, MapLibre, Nominatim.")
    doc.add_paragraph("Cas d'utilisation: s'inscrire, rechercher, envoyer lead, chatter, réserver, aftercare, stencils.")

    doc.add_heading('2.3 Spécifications fonctionnelles', level=2)
    doc.add_paragraph("Flux principaux: création profil public (géocodage adresse), découverte sur carte, messagerie temps réel, acceptation lead + aftercare.")

    doc.add_heading('2.4 Technologies utilisées', level=2)
    doc.add_paragraph("Web: Next.js 14, TypeScript, Tailwind, MapLibre GL JS.")
    doc.add_paragraph("iOS: SwiftUI, CoreLocation, PhotosUI, Combine.")
    doc.add_paragraph("Backend: Firebase Auth, RTDB, Storage; Nominatim.")
    doc.add_paragraph("Outils: GitHub, VS Code, Xcode, Firebase Console.")

    doc.add_heading('2.5 Exigences logiques de bases de données', level=2)
    doc.add_paragraph("Base NoSQL RTDB dénormalisée, hiérarchie peu profonde, duplication contrôlée.")
    # Table: Structure DB (simplified)
    add_table(
        doc,
        headers=["Espace Firebase", "Clé(s)", "Champs majeurs", "Type", "Description"],
        rows=[
            ["publicProfiles/{uid}", "uid", "role, displayName, city, address, styles, coverURL, latitude, longitude, createdAt, updatedAt", "Object", "Profils publics artistes"],
            ["leads/{leadId}", "leadId", "clientId, artistId, message, status, aftercareId, createdAt, updatedAt", "Object", "Demandes de tatouage"],
            ["aftercareByClient/{clientId}/{aftercareId}", "clientId, aftercareId", "title, steps[], status, assignedAt, artistId, completedAt", "Object", "Plans indexés par client"],
            ["aftercareByArtist/{artistId}/{aftercareId}", "artistId, aftercareId", "title, steps[], status, assignedAt, clientId, completedAt", "Object", "Plans indexés par artiste"],
            ["threads/{threadId}", "threadId", "participants[], lastMessage, lastMessageTimestamp, updatedAt", "Object", "Conversations"],
            ["threads/{threadId}/messages/{messageId}", "threadId, messageId", "senderId, text, timestamp, read", "Object", "Messages"],
            ["bookings/{bookingId}", "bookingId", "clientId, artistId, date, time, duration, status, createdAt, notes", "Object", "Réservations"],
            ["stencils/{artistId}/{stencilId}", "artistId, stencilId", "imageURL, title, description, uploadedAt, associatedLeadId", "Object", "Modèles/stencils"],
        ],
    )

    doc.add_heading('2.6 Architecture du logiciel', level=2)
    doc.add_paragraph("Architecture matérielle et logicielle (Figures 11-12) avec clients Web/iOS, Firebase (Auth, RTDB, Storage), services externes Nominatim/MapLibre.")

    # 3. DESCRIPTION DÉTAILLÉE (selected highlights)
    doc.add_heading('3. DESCRIPTION DÉTAILLÉE', level=1)
    doc.add_heading('3.1 Interfaces externes', level=2)
    doc.add_paragraph("Transactions: Auth, création profil (validation + géocodage + upload cover), recherche, envoi de message, gestion de lead.")

    doc.add_heading('3.4 Exigences logiques de bases de données', level=2)
    # Big two-column table summarizing logical requirements
    add_table(
        doc,
        headers=["Dimension", "Description détaillée"],
        rows=[
            ["Types d'informations", "Profils, Leads, Aftercare, Threads, Messages, Bookings, Stencils"],
            ["Fréquence d'utilisation", "Lectures fréquentes: profils/messages; Écritures fréquentes: messages"],
            ["Capacité d'accès", "Temps réel via WebSocket; accès public/privé selon règles; concurrence gérée"],
            ["Entités et relations", "Voir diagramme de classes; relations 1:N via IDs; duplication contrôlée"],
            ["Contraintes d'intégrité", "Règles RTDB, validations client, enums restreints, timestamps serveur, lat/lng validés"],
            ["Exigences de rétention", "Conservation durable des données clés; backups Firebase"],
            ["Exigences de stockage", "Estimations tailles; images compressées; scalabilité Firebase"],
        ],
    )

    doc.add_heading('Règles de sécurité Firebase (extrait)', level=3)
    firebase_rules = """{
  "rules": {
    "publicProfiles": {
      "$uid": {
        ".read": "auth != null",
        ".write": "auth != null && auth.uid == $uid"
      }
    },
    "leads": {
      "$leadId": {
        ".read": "auth != null && (data.child('clientId').val() == auth.uid || data.child('artistId').val() == auth.uid)",
        ".write": "auth != null"
      }
    },
    "threads": {
      "$threadId": {
        ".read": "auth != null && data.child('participants').val().contains(auth.uid)",
        ".write": "auth != null && data.child('participants').val().contains(auth.uid)"
      }
    },
    "aftercareByClient": {
      "$clientId": {
        ".read": "auth.uid == $clientId",
        ".write": "false"
      }
    },
    "aftercareByArtist": {
      "$artistId": {
        ".read": "auth.uid == $artistId",
        ".write": "auth.uid == $artistId"
      }
    }
  }
}"""
    add_code_block(doc, firebase_rules)

    doc.add_heading('3.5 Contraintes de conception', level=2)
    doc.add_paragraph("Standards code (TS/Swift), normes Firebase, contraintes réseau/quotas, conformité RGPD/PIPEDA, accessibilité WCAG.")

    doc.add_heading('3.6 Exigences non-fonctionnelles', level=2)
    add_table(
        doc,
        headers=["Caractéristique", "Exigence détaillée"],
        rows=[
            ["Fiabilité", "Gestion d'erreurs, retry, validations, logs, 99.5% cible"],
            ["Disponibilité", "SLA Firebase 99.95%, tolérance pannes, fallbacks"],
            ["Sécurité", "Auth JWT, règles strictes, TLS 1.3, URLs signées"],
            ["Entretien", "Code modulaire, typage fort, tests, doc, CI/CD"],
            ["Modularité", "Libs dédiées (aftercare, profiles, chat, auth, leads)"],
            ["Performance", "Lazy loading, pagination, cache, <3s, <500ms"],
            ["Portabilité", "Web navigateurs modernes, iOS 15+, responsive"],
            ["Utilisabilité", "UX intuitive, accessibilité, aide contextuelle"],
            ["Scalabilité", "10k+ utilisateurs, RTDB scalable, CDN"],
        ],
    )

    # 4. INFORMATIONS COMPLÉMENTAIRES (summarized)
    doc.add_heading('4. INFORMATIONS COMPLÉMENTAIRES', level=1)
    doc.add_heading('4.1 Index', level=2)
    doc.add_paragraph("Aftercare, API, Artiste, Auth, Booking, Client, Firebase, Géocodage, Lead, MapLibre, Message, Next.js, Nominatim, Profil public, RTDB, Stencil, SwiftUI, Thread, TypeScript")

    doc.add_heading('4.2 Annexes (sélection)', level=2)
    doc.add_paragraph("Annexes UC-01 à UC-07 détaillant les cas d'utilisation majeurs (messagerie, recherche, enregistrement, aftercare, booking).")

    doc.add_heading('CONCLUSION', level=1)
    doc.add_paragraph(
        "Ce rapport présente InkMatching dans son ensemble : architecture moderne (Next.js/SwiftUI/Firebase), fonctionnalités complètes (découverte, chat, leads, booking, aftercare), sécurité renforcée, UX soignée et documentation exhaustive. Le système est prêt au déploiement et extensible (paiements, notation, Android, API publique)."
    )

    # Risques et mitigation (table)
    doc.add_heading('Annexe – Risques et mitigation', level=2)
    add_table(
        doc,
        headers=["Risque", "Probabilité", "Impact", "Mitigation"],
        rows=[
            ["Dépassement quota Nominatim", "Moyenne", "Moyen", "Cache, limitation requêtes, service backup"],
            ["Quota Firebase dépassé", "Faible", "Élevé", "Surveillance, optimisation, alertes"],
            ["Données sensibles exposées", "Faible", "Critique", "Règles strictes, audits, validation entrées"],
            ["Performance dégradée", "Moyenne", "Moyen", "Lazy loading, pagination, optimisation images"],
        ],
    )


def save_docx_and_pdf():
    doc = Document()
    set_styles(doc)
    add_cover_page(doc)
    add_toc(doc)
    add_footer_page_numbers(doc)
    add_content(doc)

    Path(DOWNLOADS).mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)

    # Optional PDF text export (simple)
    text_for_pdf = (
        "RAPPORT FINAL – PROJET SYNTHÈSE\n\n" \
        "Application: InkMatching — Collège Gérald-Godin\n\n" \
        "Ce PDF est une version texte simplifiée. La version DOCX contient la mise en forme complète (couverture, TOC, en-têtes/pieds, tableaux).\n\n"
    )
    if REPORTLAB_AVAILABLE:
        c = canvas.Canvas(OUT_PDF, pagesize=LETTER)
        width, height = LETTER
        x = inch
        y = height - inch
        for line in text_for_pdf.splitlines():
            c.drawString(x, y, line)
            y -= 14
            if y < inch:
                c.showPage()
                y = height - inch
        c.showPage()
        c.save()


if __name__ == '__main__':
    save_docx_and_pdf()
    print(f"OK: DOCX -> {OUT_DOCX}")
    if REPORTLAB_AVAILABLE:
        print(f"OK: PDF  -> {OUT_PDF}")
    else:
        print("Note: reportlab non disponible, PDF non généré.")
